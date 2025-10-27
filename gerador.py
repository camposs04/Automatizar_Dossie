from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches
import pypandoc
import os

# === Caminhos ===
CAMINHO_IMAGEM1 = "/home/camposs/Desktop/estagio/desafio/balanco_pt1.png"
CAMINHO_IMAGEM2 = "/home/camposs/Desktop/estagio/desafio/balanco_pt2.png"
CAMINHO_IMAGEM3 = "/home/camposs/Desktop/estagio/desafio/dem_result.png"

CAMINHO_TEMPLETE = "/home/camposs/Desktop/estagio/desafio/templete_base_ofc.docx"
CAMINHO_NOVO_DOC = "/home/camposs/Desktop/estagio/desafio/Document.docx"

CAMINHO_BASE_FORM_MD = "/home/camposs/Desktop/estagio/desafio/base_form.md"
CAMINHO_CART_MD = "/home/camposs/Desktop/estagio/desafio/cart_responsb.md"

TEMP_BASE_DOCX = "/home/camposs/Desktop/estagio/desafio/temporario_base.docx"
TEMP_CART_DOCX = "/home/camposs/Desktop/estagio/desafio/temp_cart.docx"
TEMP_RENDERED = "/home/camposs/Desktop/estagio/desafio/temp_rendered.docx"


pypandoc.convert_file(CAMINHO_BASE_FORM_MD, 'docx', outputfile=TEMP_BASE_DOCX)
pypandoc.convert_file(CAMINHO_CART_MD, 'docx', outputfile=TEMP_CART_DOCX)


doc = DocxTemplate(CAMINHO_TEMPLETE)

# === Inserir imagens ===
balanco_pt1_img = InlineImage(doc, CAMINHO_IMAGEM1, width=Inches(6))
balanco_pt2_img = InlineImage(doc, CAMINHO_IMAGEM2, width=Inches(6))
demstr_result_img = InlineImage(doc, CAMINHO_IMAGEM3, width=Inches(6))

context = {
    'nome_empresa': 'TESTE do Nome',
    'periodo_anual': 'Junho a Agosto 2030',
    'cnpj_empresa': '23.766.826/0001-61',
    'data_dem_encerradas': '07/02/2005',
    'razao_social_empresa': 'TESTE RAZAO EMPRESA LTDA SCP VTG',
    'periodo_em_data': '07 a 12/2030',
    'balanco_patrimonial_pt1': balanco_pt1_img,
    'balanco_patrimonial_pt2': balanco_pt2_img,
    'demontr_resultado': demstr_result_img,
    'nome_socio1_adm': 'Nome TesTE 1',
    'nome_socio2_adm': 'Nome TesTE 2',
    'cpf_socio1_adm': '089.038.947-98',
    'cpf_socio2_adm': '089.128.947-97',
    'explic_demonstr': '[[EXP_DEMONSTR]]',
    'carta_responsb': '[[CARTA_RESP]]'
}

doc.render(context)
doc.save(TEMP_RENDERED)


def insert_docx_at_placeholder(main_doc: Document, placeholder: str, insert_doc_path: str):
    found = False
    insert_doc = Document(insert_doc_path)

    for paragraph in main_doc.paragraphs:
        if placeholder in paragraph.text:
            found = True
            print(f"➡ Inserindo conteúdo em: {placeholder}")

            paragraph.text = paragraph.text.replace(placeholder, "")

            for element in reversed(insert_doc.element.body):
                paragraph._element.addnext(element)

    if not found:
        print(f"⚠️ Marcador '{placeholder}' não encontrado no documento!")

final_doc = Document(TEMP_RENDERED)


insert_docx_at_placeholder(final_doc, '[[EXP_DEMONSTR]]', TEMP_BASE_DOCX)
insert_docx_at_placeholder(final_doc, '[[CARTA_RESP]]', TEMP_CART_DOCX)

final_doc.save(CAMINHO_NOVO_DOC)

for temp_file in [TEMP_BASE_DOCX, TEMP_CART_DOCX, TEMP_RENDERED]:
    try:
        os.remove(temp_file)
    except OSError as e:
        print(f"Erro ao remover {temp_file}: {e}")

print("✅ Documento final gerado com sucesso!")
print(f"📄 Caminho: {CAMINHO_NOVO_DOC}")
