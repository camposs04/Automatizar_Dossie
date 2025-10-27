import streamlit as st
import os
import tempfile
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches
from io import BytesIO
# Restaurando o pypandoc para conversão de Markdown
import pypandoc

# =========================================================================
# === CORREÇÃO CRÍTICA PARA DEPLOY (STREAMLIT CLOUD / PANDOC) ===========
# Força o pypandoc a baixar o binário do Pandoc se não for encontrado.
# Mantemos esta lógica, pois ela é o último recurso para garantir o Pandoc.
# =========================================================================
if 'pandoc_checked' not in st.session_state:
    st.session_state['pandoc_checked'] = False

if not st.session_state['pandoc_checked']:
    try:
        pypandoc.get_pandoc_path()
        st.session_state['pandoc_checked'] = True
    except Exception:
        # Se não encontrar, tenta baixar o binário localmente
        st.warning("Pandoc não encontrado no PATH. Tentando download automático do binário (processo único, pode levar alguns segundos).")
        try:
            pypandoc.download_pandoc()
            st.session_state['pandoc_checked'] = True
            st.success("Pandoc configurado com sucesso! Clique no botão 'GERAR DOCUMENTO FINAL' novamente.")
        except Exception as download_e:
            st.error(f"Falha ao baixar Pandoc automaticamente. Erro: {download_e}")


# --- FUNÇÃO AUXILIAR: Inserir DOCX no Placeholder (Restaurada) ---
def insert_docx_at_placeholder(main_doc: Document, placeholder: str, insert_doc_path: str):
    """Substitui marcador por outro documento DOCX (mantendo a ordem correta)."""
    insert_doc = Document(insert_doc_path)
    for paragraph in main_doc.paragraphs:
        if placeholder in paragraph.text:
            # Remove o marcador do texto
            paragraph.text = paragraph.text.replace(placeholder, "")
            # Insere conteúdo do documento externo (reversed para manter a ordem)
            for element in reversed(insert_doc.element.body):
                paragraph._element.addnext(element)
            return True


# --- 1. Lógica de Geração do Documento ---
def generate_document(input_data):
    temp_paths = {}
    
    # === A) Preparar Caminhos Temporários para Imagens e MDs ===
    
    for key, uploaded_file in input_data['uploads'].items():
        if uploaded_file is not None:
            suffix = os.path.splitext(uploaded_file.name)[1]

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                temp_paths[key] = tmp_file.name
        else:
            return None, f"O arquivo {key} é obrigatório!"
    
    # === B) Definir os Caminhos Finais para a Lógica de Geração ===
    CAMINHO_TEMPLETE = "templete_base_ofc.docx"
    
    TEMP_BASE_DOCX = os.path.join(tempfile.gettempdir(), "temporario_base.docx")
    TEMP_CART_DOCX = os.path.join(tempfile.gettempdir(), "temp_cart.docx")
    TEMP_RENDERED = os.path.join(tempfile.gettempdir(), "temp_rendered.docx")
    
    final_docx_buffer = BytesIO() 

    try:
        # 1. CONVERSÃO DE MD PARA DOCX (Restaurada)
        # O pypandoc converte o Markdown para um arquivo DOCX temporário
        pypandoc.convert_file(temp_paths['explic_demonstr_file'], 'docx', outputfile = TEMP_BASE_DOCX)
        pypandoc.convert_file(temp_paths['carta_responsb_file'], 'docx', outputfile = TEMP_CART_DOCX)
        
        doc = DocxTemplate(CAMINHO_TEMPLETE)

        balanco_pt1_img = InlineImage(doc, temp_paths['balanco_pt1_file'], width=Inches(6))
        balanco_pt2_img = InlineImage(doc, temp_paths['balanco_pt2_file'], width=Inches(6))
        demstr_result_img = InlineImage(doc, temp_paths['demstr_result_file'], width=Inches(6))

        context = {
            'nome_empresa': input_data['nome_empresa'],
            'periodo_anual': input_data['periodo_anual'],
            'cnpj_empresa': input_data['cnpj_empresa'],
            'data_dem_encerradas': input_data['data_dem_encerradas'],
            'razao_social_empresa': input_data['razao_social_empresa'],
            'periodo_em_data': input_data['periodo_em_data'],
            'balanco_patrimonial_pt1': balanco_pt1_img,
            'balanco_patrimonial_pt2': balanco_pt2_img,
            'demontr_resultado': demstr_result_img,
            'nome_socio1': input_data['nome_socio1'],
            'nome_socio2': input_data['nome_socio2'],
            'cargo_socio1': input_data['cargo_socio1'],
            'cargo_socio2': input_data['cargo_socio2'],
            'cpf_socio1': input_data['cpf_socio1'],
            'cpf_socio2': input_data['cpf_socio2'],
            'explic_demonstr': '[[EXP_DEMONSTR]]', 
            'carta_responsb': '[[CARTA_RESP]]'
        }

        doc.render(context)
        doc.save(TEMP_RENDERED)

        final_doc = Document(TEMP_RENDERED)

        # 3. INSERÇÃO DOS DOCX (gerados pelo Pandoc)
        insert_docx_at_placeholder(final_doc, '[[EXP_DEMONSTR]]', TEMP_BASE_DOCX)
        insert_docx_at_placeholder(final_doc, '[[CARTA_RESP]]', TEMP_CART_DOCX)


        final_doc.save(final_docx_buffer)
        final_docx_buffer.seek(0)
        
        return final_docx_buffer.getvalue(), None

    except Exception as e:
        # Tratamento de erro detalhado para deploy
        if "No such file or directory" in str(e) and CAMINHO_TEMPLETE in str(e):
            return None, f"Erro: O template DOCX '{CAMINHO_TEMPLETE}' não foi encontrado no repositório. Certifique-se de que ele foi enviado ao GitHub."
        if "No pandoc was found" in str(e):
             return None, f"Erro: O Pandoc é necessário para converter Markdown. O download automático falhou. Erro detalhado: {e}"
        
        return None, f"Erro durante a geração: {e}"
    
    finally:
        temp_files_to_clean = [
            TEMP_BASE_DOCX, TEMP_CART_DOCX, TEMP_RENDERED, 
            temp_paths.get('balanco_pt1_file'), temp_paths.get('balanco_pt2_file'), 
            temp_paths.get('demstr_result_file'), temp_paths.get('explic_demonstr_file'), 
            temp_paths.get('carta_responsb_file')
        ]
        
        for temp_file in [f for f in temp_files_to_clean if f]:
            try:
                os.remove(temp_file)
            except Exception:
                pass


# --- 2. Interface Streamlit ---

st.set_page_config(page_title="Gerador de Demonstrações Contábeis", layout="wide")
st.title("📄 Gerador Automático de Documentos Contábeis")
st.markdown("Preencha os campos e faça o upload dos arquivos para gerar o dossiê final.")

tab1, tab2, tab3 = st.tabs(["Dados da Empresa/Períodos", "Dados dos Sócios", "Upload de Arquivos"])

input_data = {}

with tab1:
    col1, col2 = st.columns(2)
    
    with col1:
        input_data['nome_empresa'] = st.text_input("Nome Fantasia da Empresa", value="TESTE do Nome")
        input_data['razao_social_empresa'] = st.text_input("Razão Social", value="TESTE RAZAO EMPRESA LTDA SCP VTG")
        input_data['cnpj_empresa'] = st.text_input("CNPJ", value="23.766.826/0001-61", help="Formato: 00.000.000/0000-00")

    with col2:
        input_data['periodo_anual'] = st.text_input("Período Anual (Descrição)", value="Junho a Agosto 2030")
        input_data['data_dem_encerradas'] = st.text_input("Demonstrações Contábeis Encerradas em", value="07/02/2005", help="Formato: DD/MM/AAAA")
        input_data['periodo_em_data'] = st.text_input("Período de Referência", value="07 a 12/2030")

with tab2:
    st.subheader("Dados dos Sócios")
    col3, col4 = st.columns(2)
    
    with col3:
        input_data['nome_socio1'] = st.text_input("Nome do Sócio 1", value="Nome TesTE 1")
        input_data['cargo_socio1'] = st.text_input("Cargo do Sócio 1", value = "Cargo Teste 1")
        input_data['cpf_socio1'] = st.text_input("CPF do Sócio 1", value="089.038.947-98")
    
    with col4:
        input_data['nome_socio2'] = st.text_input("Nome do Sócio 2", value="Nome TesTE 2")
        input_data['cargo_socio2'] = st.text_input("Cargo do Sócio 2", value = "Cargo Teste 2")
        input_data['cpf_socio2'] = st.text_input("CPF do Sócio 2", value="089.128.947-97")

input_data['uploads'] = {}

with tab3:
    st.subheader("Imagens de Demonstrações (PNG ou JPG)")
    col5, col6 = st.columns(2)
    
    with col5:
        input_data['uploads']['balanco_pt1_file'] = st.file_uploader("Balanco Patrimonial (Ativo)", type=["png", "jpg"], key='balanco_1')
        input_data['uploads']['balanco_pt2_file'] = st.file_uploader("Balanco Patrimonial (Passivo)", type=["png", "jpg"], key='balanco_2')
    with col6:
        input_data['uploads']['demstr_result_file'] = st.file_uploader("Demonstração do Resultado (DRE)", type=["png", "jpg"], key='dre')

    st.subheader("Arquivos de Texto (Markdown - Recomendado)")
    col7, col8 = st.columns(2)
    with col7:
        # Tipo de arquivo ALTERADO para .md
        input_data['uploads']['explic_demonstr_file'] = st.file_uploader("Notas Explicativas", type=["md"], key='notas')
    with col8:
        # Tipo de arquivo ALTERADO para .md
        input_data['uploads']['carta_responsb_file'] = st.file_uploader("Carta de Responsabilidade", type=["md"], key='carta')

if st.button("✅ GERAR DOCUMENTO FINAL", type="primary"):
    
    if not st.session_state['pandoc_checked']:
        st.warning("Aguarde a mensagem de sucesso do Pandoc e clique aqui novamente.")
    else:
        required_files = [
            'balanco_pt1_file', 'balanco_pt2_file', 'demstr_result_file', 
            'explic_demonstr_file', 'carta_responsb_file'
        ]
        
        all_files_uploaded = all(input_data['uploads'][f] is not None for f in required_files)
        
        if all_files_uploaded:
            with st.spinner("Gerando documento... Isso pode levar alguns segundos."):
                file_data, error = generate_document(input_data)
            
            if file_data:
                st.success("Documento gerado com sucesso!")
                st.download_button(
                    label="Clique para Baixar Document.docx",
                    data=file_data,
                    file_name=f"Dossie_Contabil_{input_data['nome_empresa']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(f"Falha na geração do documento. Detalhes: {error}")
                
        else:
            st.warning("Por favor, faça o upload de todos os 5 arquivos antes de gerar.")
