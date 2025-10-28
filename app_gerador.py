import streamlit as st
import os
import tempfile
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches
from io import BytesIO
import fitz
import datetime

def pdf_balanco_duas_paginas(pdf_path):
    doc = fitz.open(pdf_path)
    images = []

    for i in range(2):
        if i < len(doc):
            pix = doc[i].get_pixmap(dpi = 200)
            img_bytes = pix.tobytes("png")
            images.append(img_bytes)
        else:
            images.append(None)
    return images

def pdf_to_images(pdf_path):
    images = []
    doc = fitz.open(pdf_path)
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        images.append(img_bytes)
    return images

def insert_pdf_at_placeholder(main_doc, placeholder, pdf_path):
    images = pdf_to_images(pdf_path)
    for paragraph in main_doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            for img in images:
                run = paragraph.add_run()
                run.add_picture(BytesIO(img), width=Inches(6))
            return True

def insert_docx_at_placeholder(main_doc: Document, placeholder: str, insert_doc_path: str):
    insert_doc = Document(insert_doc_path)
    for paragraph in main_doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            for element in reversed(insert_doc.element.body):
                paragraph._element.addnext(element)
            return True

def generate_document(input_data):
    temp_paths = {}
    
    # === A) Preparar Caminhos Temporários para Imagens ===
    
    for key, uploaded_file in input_data['uploads'].items():
        if uploaded_file is not None:
            suffix = os.path.splitext(uploaded_file.name)[1]

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                temp_paths[key] = tmp_file.name
        else:
            return None, f"O arquivo {key} é obrigatório!"
    
    # === B) Definir os Caminhos Finais para a Lógica de Geração ===
    CAMINHO_TEMPLETE = "templete_base_ofc.docx"
    
    TEMP_BASE_DOCX = os.path.join(tempfile.gettempdir(), "temporario_base.docx")
    TEMP_CART_DOCX = os.path.join(tempfile.gettempdir(), "temp_cart.docx")
    TEMP_RENDERED = os.path.join(tempfile.gettempdir(), "temp_rendered.docx")
    
    final_docx_buffer = BytesIO() 

    try:
        
        doc = DocxTemplate(CAMINHO_TEMPLETE)

        balanco_imgs = pdf_balanco_duas_paginas(temp_paths['balanco_file'])
        balanco_pt1_img = InlineImage(doc, BytesIO(balanco_imgs[0]), width=Inches(6))
        balanco_pt2_img = InlineImage(doc, BytesIO(balanco_imgs[1]), width=Inches(6))

        context = {
            'nome_empresa': input_data['nome_empresa'],
            'periodo_anual': input_data['periodo_anual'],
            'cnpj_empresa': input_data['cnpj_empresa'],
            'data_dem_encerradas': input_data['data_dem_encerradas'],
            'razao_social_empresa': input_data['razao_social_empresa'],
            'periodo_em_data': input_data['periodo_em_data'],
            'balanco_patrimonial_pt1': balanco_pt1_img,
            'balanco_patrimonial_pt2': balanco_pt2_img,
            'demontr_resultado': '[[DEMONSTR_RESULTADO]]',
            'socios': input_data['socios'],
            'explic_demonstr': '[[EXP_DEMONSTR]]', 
            'carta_responsb': '[[CARTA_RESP]]'
        }

        doc.render(context)
        doc.save(TEMP_RENDERED)

        final_doc = Document(TEMP_RENDERED)

        insert_pdf_at_placeholder(final_doc, '[[DEMONSTR_RESULTADO]]', temp_paths['demstr_result_file'])
        insert_docx_at_placeholder(final_doc, '[[EXP_DEMONSTR]]', temp_paths['explic_demonstr_file'])
        insert_docx_at_placeholder(final_doc, '[[CARTA_RESP]]', temp_paths['carta_responsb_file'])

        final_doc.save(final_docx_buffer)
        final_docx_buffer.seek(0)
        
        return final_docx_buffer.getvalue(), None

    except Exception as e:
        if "No such file or directory" in str(e) and CAMINHO_TEMPLETE in str(e):
            return None, f"Erro: O template DOCX '{CAMINHO_TEMPLETE}' não foi encontrado no repositório. Certifique-se de que ele foi enviado ao GitHub."
        
        if "No pandoc was found" in str(e):
             return None, f"Erro: O Pandoc é necessário para converter Markdown. Por favor, instale o Pandoc no ambiente ou use uma solução de deploy que o inclua. Erro detalhado: {e}"

        if "index out of range" in str(e):
             return None, f"Erro: O arquivo 'Balanco Patrimonial' (PDF) deve ter pelo menos 2 páginas. Detalhes: {e}"
       
        return None, f"Erro durante a geração: {e}"
    
    finally:
        temp_files_to_clean = [
            TEMP_BASE_DOCX, TEMP_CART_DOCX, TEMP_RENDERED, 
            temp_paths.get('balanco_pt1_file'), temp_paths.get('balanco_pt2_file'), 
            temp_paths.get('demstr_result_file'), temp_paths.get('explic_demonstr_file'), 
            temp_paths.get('carta_responsb_file')
        ]
        
        for temp_file in [f for f in temp_files_to_clean if f]:
            try:
                os.remove(temp_file)
            except Exception:
                pass


# --- 2. Interface Streamlit ---

st.set_page_config(page_title="Gerador de Demonstrações Contábeis", layout="wide")
st.title("📄 Gerador Automático de Documentos Contábeis")
st.markdown("Preencha os campos e faça o upload dos arquivos para gerar o dossiê final.")

tab1, tab2, tab3 = st.tabs(["Dados da Empresa/Períodos", "Dados dos Sócios", "Upload de Arquivos"])

input_data = {}

meses_pt = {
    1:"Janeiro", 2:"Fevereiro", 3:"Março", 4:"Abril",
    5:"Maio", 6:"Junho", 7:"Julho", 8:"Agosto",
    9:"Setembro", 10:"Outubro",11:"Novembro",12:"Dezembro"
}

with tab1:
    col1, col2 = st.columns(2)
    
    with col1:
        input_data['nome_empresa'] = st.text_input("Nome Fantasia da Empresa", value="TESTE do Nome")
        input_data['razao_social_empresa'] = st.text_input("Razão Social", value="TESTE RAZAO EMPRESA LTDA SCP VTG")
        input_data['cnpj_empresa'] = st.text_input("CNPJ", value="23.766.826/0001-61", help="Formato: 00.000.000/0000-00")

    with col2:

        st.markdown("##### Período de Referência Contábil")

        data_inicio_default = datetime.date(datetime.date.today().year, 1,1)
        data_fim_default = datetime.date(datetime.date.today().year, 12, 31)

        data_inicio = st.date_input("Data de Início", value = data_inicio_default, key = 'data_inicio')
        data_fim = st.date_input("Data de Fim", value = data_fim_default, key = 'data_fim')
        #input_data['data_dem_encerradas'] = st.text_input("Demonstrações Contábeis Encerradas em", value="07/02/2005", help="Formato: DD/MM/AAAA")
        #input_data['periodo_em_data'] = st.text_input("Período de Referência", value="07 a 12/2030")

        mes_inicio_curto = str(data_inicio.month).zfill(2)
        ano_inicio_curto = str(data_inicio.year)[-2:]
        mes_fim_curto = str(data_fim.month).zfill(2)
        ano_fim_curto = str(data_fim.year)[-2:]
        
        if data_inicio.year != data_fim.year:
             input_data['periodo_em_data'] = f"{mes_inicio_curto}/{ano_inicio_curto} a {mes_fim_curto}/{ano_fim_curto}"
        else:
             input_data['periodo_em_data'] = f"{mes_inicio_curto} a {mes_fim_curto}/{ano_fim_curto}"
        
        mes_desc_inicio = meses_pt.get(data_inicio.month)
        mes_desc_fim = meses_pt.get(data_fim.month)
        
        if data_inicio.year == data_fim.year:
            periodo_anual_desc = f"{mes_desc_inicio} a {mes_desc_fim} de {data_inicio.year}"
        else:
            periodo_anual_desc = f"{mes_desc_inicio} de {data_inicio.year} a {mes_desc_fim} de {data_fim.year}"

        input_data['periodo_anual'] = periodo_anual_desc

        input_data['data_dem_encerradas'] = data_fim.strftime("%d/%m/%Y")
        
        st.markdown("---")
        st.markdown(f"**Período de Referência (periodo_em_data):** `{input_data['periodo_em_data']}`")
        st.markdown(f"**Descrição Anual (periodo_anual):** `{input_data['periodo_anual']}`")


with tab2:
    st.subheader("Dados dos Sócios")
    if "socios" not in st.session_state:
        st.session_state.socios = [{"nome": "", "cpf": "", "cargo": ""}]

    for i, socio in enumerate(st.session_state.socios):
        st.write(f"--- Sócio {i+1} ---")
        socio["nome"] = st.text_input(f"Nome do Sócio {i+1}", value=socio["nome"], key=f"nome_{i}")
        socio["cpf"] = st.text_input(f"CPF do Sócio {i+1}", value=socio["cpf"], key=f"cpf_{i}")
        socio["cargo"] = st.text_input(f"Cargo do Sócio {i+1}", value=socio["cargo"], key=f"cargo_{i}")
        if st.button(f"Remover Sócio {i+1}", key=f"remove_{i}"):
            st.session_state.socios.pop(i)
            st.experimental_rerun()

    if st.button("➕ Adicionar Sócio"):
        st.session_state.socios.append({"nome": "", "cpf": "", "cargo": ""})

input_data["socios"] = st.session_state.socios

input_data['uploads'] = {}

with tab3:
    st.subheader("Balancos e Demonstrações (PDF)")
    col5, col6 = st.columns(2)
    
    with col5:
        input_data['uploads']['balanco_file'] = st.file_uploader("Balanco Patrimonial (PDF)", type=["pdf"])
    with col6:
        input_data['uploads']['demstr_result_file'] = st.file_uploader("Demonstração do Resultado (DRE)", type=["pdf"])

    st.subheader("Arquivos de Texto (WORD)")
    col7, col8 = st.columns(2)
    with col7:
        input_data['uploads']['explic_demonstr_file'] = st.file_uploader("Notas Explicativas", type=["docx"], key='notas')
    with col8:
        input_data['uploads']['carta_responsb_file'] = st.file_uploader("Carta de Responsabilidade", type=["docx"], key='carta')

if st.button("✅ GERAR DOCUMENTO FINAL", type="primary"):
    required_files = [
        'balanco_file', 'demstr_result_file', 
        'explic_demonstr_file', 'carta_responsb_file'
    ]
    
    all_files_uploaded = all(input_data['uploads'][f] is not None for f in required_files)
    
    if all_files_uploaded:
        with st.spinner("Gerando documento... Isso pode levar alguns segundos."):
            file_data, error = generate_document(input_data)
        
        if file_data:
            st.success("Documento gerado com sucesso!")
            st.download_button(
                label="Clique para Baixar Document.docx",
                data=file_data,
                file_name=f"Dossie_Contabil_{input_data['nome_empresa']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error(f"Falha na geração do documento. Detalhes: {error}")
            
    else:
        st.warning("Por favor, faça o upload de todos os 5 arquivos antes de gerar.")